from flask import Blueprint, jsonify, request, Response, send_file, make_response
import pdfkit
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import *
from db import *
from utils import *

query_blueprint = Blueprint("query", __name__)


def getCourse(teacherNo, startYear, endYear):
    results = (
        db.session.query(Course, TeachCourse)
        .join(Course, TeachCourse.courseNo == Course.No)
        .filter(TeachCourse.teacherNo == teacherNo)
        .filter(TeachCourse.year >= startYear)
        .filter(TeachCourse.year <= endYear)
        .all()
    )
    courses = []
    for result in results:
        courses.append(
            {
                "No": result.Course.No,
                "name": result.Course.name,
                "takeCreditHour": result.TeachCourse.takeCreditHour,
                "year": result.TeachCourse.year,
                "semester": result.TeachCourse.semester,
            }
        )
    return courses


def getPaper(teacherNo, startYear, endYear):
    results = (
        db.session.query(Paper, PublishedPapers)
        .join(Paper, PublishedPapers.paperNo == Paper.No)
        .filter(PublishedPapers.teacherNo == teacherNo)
        .filter(Paper.publishYear >= startYear)
        .filter(Paper.publishYear <= endYear)
        .all()
    )
    papers = []
    for result in results:
        papers.append(
            {
                "No": result.Paper.No,
                "name": result.Paper.name,
                "source": result.Paper.source,
                "year": result.Paper.publishYear,
                "type": result.Paper.type,
                "level": result.Paper.level,
                "rank": result.PublishedPapers.rank,
                "isCoAuthor": result.PublishedPapers.isCoAuthor,
            }
        )
    return papers


def getProject(teacherNo, startYear, endYear):
    results = (
        db.session.query(Project, TakeProject)
        .join(Project, TakeProject.projectNo == Project.No)
        .filter(TakeProject.teacherNo == teacherNo)
        .filter(Project.startYear <= endYear)
        .filter(Project.endYear >= startYear)
        .all()
    )
    projects = []
    for result in results:
        projects.append(
            {
                "No": result.Project.No,
                "name": result.Project.name,
                "source": result.Project.source,
                "type": result.Project.type,
                "startYear": result.Project.startYear,
                "endYear": result.Project.endYear,
                "funds": result.Project.funds,
                "takeFunds": result.TakeProject.takeFunds,
            }
        )
    return projects


@query_blueprint.route("/query/<teacherNo>/<startYear>/<endYear>", methods=["GET"])
def query(teacherNo, startYear, endYear, local=False):
    response = {"status": True}
    teacher = Teacher.query.filter_by(No=teacherNo).first()
    if not teacher:
        response["status"] = False
        response["message"] = "查询失败：教师工号不存在！"
        return jsonify(response)
    response["teacher"] = teacher.json()

    try:
        response["courses"] = getCourse(teacherNo, startYear, endYear)
    except Exception as e:
        print(e)
        response["status"] = False
        response["message"] = "教学信息获取出错"
        return jsonify(response)

    try:
        response["papers"] = getPaper(teacherNo, startYear, endYear)
    except Exception as e:
        print(e)
        response["status"] = False
        response["message"] = "发表论文信息获取出错"
        return jsonify(response)

    try:
        response["projects"] = getProject(teacherNo, startYear, endYear)
    except Exception as e:
        print(e)
        response["status"] = False
        response["message"] = "承担项目信息获取出错"
        return jsonify(response)

    response["message"] = "查询成功！"
    if local:
        return (
            response["teacher"],
            response["courses"],
            response["papers"],
            response["projects"],
        )
    return jsonify(response)


def generatePDF(teacherNo, startYear, endYear):
    teacher, courses, papers, projects = query(
        teacherNo, startYear, endYear, local=True
    )
    html = """
    <html>
    <head>
    <style>
    h1 {
        text-align: center;
    }
    h2 {
        border-bottom: 1px solid black;
        padding-bottom: 0.25em;
        margin-bottom: 1.0em;
    }
    table {
        border-collapse: collapse;
        margin: 0 auto;
    }
    th, td {
        border: 1px solid black;
        padding: 0.5em;
        text-align: center;
    }
    th {
        background-color: #eee;
    }
    </style>
    </head>
    <body>
    <h1>教师教学科研工作统计（"""

    html += str(startYear) + "-" + str(endYear)

    html += """）</h1>

    <h2>教师基本信息</h2>
    <p style="font-size: 24px;">"""

    k = 6
    html += "工号：" + teacher["No"] + "&nbsp;" * k
    html += "姓名：" + teacher["name"] + "&nbsp;" * k
    html += "性别：" + genderMap[teacher["gender"]] + "&nbsp;" * k
    html += "职称：" + titleMap[teacher["title"]] + "</p>"

    html += """

    <h2>教学情况</h2>
    <table>
        <tr>
            <th>课程号</th>
            <th>课程名</th>
            <th>主讲学时</th>
            <th>学期</th>
        </tr>
"""
    for row in courses:
        html += "<tr>"
        html += "<td>" + row["No"] + "</td>"
        html += "<td>" + row["name"] + "</td>"
        html += "<td>" + str(row["takeCreditHour"]) + "</td>"
        html += "<td>" + str(row["year"]) + semesterMap[row["semester"]] + "</td>"
        html += "</tr>"

    html += """
    </table>

    <h2>发表论文情况</h2>
    <table>
        <tr>
            <th style="white-space: nowrap">论文标题</th>
            <th style="white-space: nowrap">发表源</th>
            <th style="white-space: nowrap">发表年份</th>
            <th style="white-space: nowrap">类型</th>
            <th style="white-space: nowrap">级别</th>
            <th style="white-space: nowrap">排名</th>
            <th style="white-space: nowrap">通讯作者</th>
        </tr>
    """

    for row in papers:
        html += "<tr>"
        html += "<td>" + row["name"] + "</td>"
        html += '<td style="white-space: nowrap">' + row["source"] + "</td>"
        html += "<td>" + str(row["year"]) + "</td>"
        html += '<td style="white-space: nowrap">' + paperTypeMap[row["type"]] + "</td>"
        html += (
            '<td style="white-space: nowrap">' + paperLevelMap[row["level"]] + "</td>"
        )
        html += "<td>第" + str(row["rank"]) + "</td>"
        html += "<td>" + coAuthorMap[row["isCoAuthor"]] + "</td>"
        html += "</tr>"

    html += """
    </table>

    <h2>承担项目情况</h2>
    <table>
        <tr>
            <th style="white-space: nowrap">项目名称</th>
            <th style="white-space: nowrap">项目来源</th>
            <th style="white-space: nowrap">项目级别</th>
            <th style="white-space: nowrap">起止年份</th>
            <th style="white-space: nowrap">总经费</th>
            <th style="white-space: nowrap">承担经费</th>
        </tr>
    """

    for row in projects:
        html += "<tr>"
        html += "<td>" + row["name"] + "</td>"
        html += "<td>" + row["source"] + "</td>"
        html += "<td>" + projectMap[row["type"]] + "</td>"
        html += "<td>" + str(row["startYear"]) + "-" + str(row["endYear"]) + "</td>"
        html += "<td>" + str(row["funds"]) + "</td>"
        html += "<td>" + str(row["takeFunds"]) + "</td>"
        html += "</tr>"

    html += """
    </table>

    </body>
    </html>
    """

    # 将 HTML 转换为 PDF 文件
    options = {
        "page-size": "A4",
        "encoding": "UTF-8",
    }
    pdfkit.from_string(html, "output.pdf", options=options)
    # time.sleep(1)


@query_blueprint.route("/export/pdf/<teacherNo>/<startYear>/<endYear>", methods=["GET"])
def export_pdf(teacherNo, startYear, endYear):
    generatePDF(teacherNo, startYear, endYear)
    # 读取PDF文件的二进制数据
    with open("output.pdf", "rb") as f:
        pdf_data = f.read()

    # 构造响应对象，并设置响应头和响应体
    response = make_response(pdf_data)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = "attachment; filename=output.pdf"

    return response


def generate_xlsx(teacherNo, startYear, endYear):
    teacher, courses, papers, projects = query(
        teacherNo, startYear, endYear, local=True
    )
    workbook = openpyxl.Workbook()

    # sheet1
    sheet1 = workbook.active
    sheet1.title = "教师"
    sheet1.append(["工号", "姓名", "性别", "职称"])
    sheet1.append(
        [
            teacher["No"],
            teacher["name"],
            genderMap[teacher["gender"]],
            titleMap[teacher["title"]],
        ]
    )

    # sheet2
    sheet2 = workbook.create_sheet(title="主讲课程")
    sheet2.append(["课程号", "课程名", "年份", "学期", "主讲学时"])
    for row in courses:
        sheet2.append(
            [
                row["No"],
                row["name"],
                row["year"],
                semesterMap[row["semester"]],
                row["takeCreditHour"],
            ]
        )

    # sheet3
    sheet3 = workbook.create_sheet(title="发表论文")
    sheet3.append(["论文序号", "论文名", "发表源", "发表年份", "类型", "级别", "排名", "是否通讯作者"])
    for row in papers:
        sheet3.append(
            [
                row["No"],
                row["name"],
                row["source"],
                row["year"],
                paperTypeMap[row["type"]],
                paperLevelMap[row["level"]],
                row["rank"],
                coAuthorMap[row["isCoAuthor"]],
            ]
        )

    # sheet4
    sheet4 = workbook.create_sheet(title="承担项目")
    sheet4.append(["项目号", "项目名称", "项目来源", "项目类型", "开始年份", "结束年份", "项目总经费", "承担经费"])
    for row in projects:
        sheet4.append(
            [
                row["No"],
                row["name"],
                row["source"],
                projectMap[row["type"]],
                row["startYear"],
                row["endYear"],
                row["funds"],
                row["takeFunds"],
            ]
        )

    # save
    workbook.save("output.xlsx")


@query_blueprint.route(
    "/export/xlsx/<teacherNo>/<startYear>/<endYear>", methods=["GET"]
)
def export_xlsx(teacherNo, startYear, endYear):
    generate_xlsx(teacherNo, startYear, endYear)

    with open("output.xlsx", "rb") as f:
        data = f.read()

    response = make_response(data)
    response.headers["Content-Disposition"] = "attachment; filename=output.xlsx"
    response.mimetype = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    return response


def generate_docx(teacherNo, startYear, endYear):
    teacher, courses, papers, projects = query(
        teacherNo, startYear, endYear, local=True
    )

    document = Document()
    title = document.add_heading(
        "教师教学科研工作统计（" + str(startYear) + "-" + str(endYear) + "）", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("教师信息", level=2)
    k = 10
    text = ""
    text += "工号：" + teacher["No"] + " " * k
    text += "姓名：" + teacher["name"] + " " * k
    text += "性别：" + genderMap[teacher["gender"]] + " " * k
    text += "职称：" + titleMap[teacher["title"]] + "\n"
    document.add_paragraph(text)

    document.add_heading("教学情况", level=2)
    table = document.add_table(rows=1, cols=4, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "课程号"
    hdr_cells[1].text = "课程名"
    hdr_cells[2].text = "主讲学时"
    hdr_cells[3].text = "学期"
    for row in courses:
        row_cells = table.add_row().cells
        row_cells[0].text = row["No"]
        row_cells[1].text = row["name"]
        row_cells[2].text = str(row["takeCreditHour"])
        row_cells[3].text = str(row["year"]) + semesterMap[row["semester"]]

    document.add_heading("发表论文情况", level=2)
    table = document.add_table(rows=1, cols=7, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "论文标题"
    hdr_cells[1].text = "发表源"
    hdr_cells[2].text = "发表年份"
    hdr_cells[3].text = "类型"
    hdr_cells[4].text = "级别"
    hdr_cells[5].text = "排名"
    hdr_cells[6].text = "通讯作者"
    for row in papers:
        row_cells = table.add_row().cells
        row_cells[0].text = row["name"]
        row_cells[1].text = row["source"]
        row_cells[2].text = str(row["year"])
        row_cells[3].text = paperTypeMap[row["type"]]
        row_cells[4].text = paperLevelMap[row["level"]]
        row_cells[5].text = str(row["rank"])
        row_cells[6].text = coAuthorMap[row["isCoAuthor"]]

    document.add_heading("承担项目情况", level=2)
    table = document.add_table(rows=1, cols=6, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "项目名称"
    hdr_cells[1].text = "项目来源"
    hdr_cells[2].text = "项目类型"
    hdr_cells[3].text = "起止年份"
    hdr_cells[4].text = "总经费"
    hdr_cells[5].text = "承担经费"
    for row in projects:
        row_cells = table.add_row().cells
        row_cells[0].text = row["name"]
        row_cells[1].text = row["source"]
        row_cells[2].text = projectMap[row["type"]]
        row_cells[3].text = str(row["startYear"]) + "-" + str(row["endYear"])
        row_cells[4].text = str(row["funds"])
        row_cells[5].text = str(row["takeFunds"])

    document.save("output.docx")


@query_blueprint.route(
    "/export/docx/<teacherNo>/<startYear>/<endYear>", methods=["GET"]
)
def export_docx(teacherNo, startYear, endYear):
    generate_docx(teacherNo, startYear, endYear)
    with open("output.docx", "rb") as f:
        data = f.read()

    response = make_response(data)
    response.headers["Content-Disposition"] = "attachment; filename=output.docx"
    response.mimetype = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    return response
